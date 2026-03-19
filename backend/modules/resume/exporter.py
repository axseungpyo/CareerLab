"""Resume exporter — generate DOCX, PDF, Markdown, TXT, HTML, JSON from resume data."""

import io
import json as json_lib

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from supabase import create_client

from config.settings import get_settings


class ResumeExporter:
    """Export resume to DOCX format."""

    def __init__(self):
        settings = get_settings()
        self._db = create_client(
            settings.supabase_url, settings.supabase_service_role_key
        )

    def to_docx(self, resume_id: str) -> bytes:
        """Generate a formatted DOCX from resume data."""
        # Load resume with items and company info
        result = (
            self._db.table("resumes")
            .select("*, resume_items(*), company_analyses(company_name)")
            .eq("id", resume_id)
            .execute()
        )
        if not result.data:
            raise ValueError("자소서를 찾을 수 없습니다.")

        resume = result.data[0]
        items = sorted(resume.get("resume_items", []), key=lambda x: x.get("created_at", ""))
        company = resume.get("company_analyses", {}).get("company_name", "")

        doc = Document()

        # Title
        title = doc.add_heading(resume.get("title", "자기소개서"), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if company:
            subtitle = doc.add_paragraph(f"지원 기업: {company}")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(11)

        doc.add_paragraph("")

        # Items
        for i, item in enumerate(items, 1):
            q = doc.add_heading(f"{i}. {item.get('question', '')}", level=2)
            q.runs[0].font.size = Pt(12)

            answer = item.get("answer", "")
            p = doc.add_paragraph(answer)
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.size = Pt(10.5)

            if item.get("char_limit"):
                info = doc.add_paragraph(
                    f"(글자수: {len(answer)}/{item['char_limit']})"
                )
                info.runs[0].font.size = Pt(9)
                info.runs[0].font.color.rgb = None

            doc.add_paragraph("")

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def to_pdf(self, resume_id: str) -> bytes:
        """Generate PDF by converting DOCX → PDF via docx2pdf or fallback HTML."""
        # Use HTML-based PDF generation for portability
        result = (
            self._db.table("resumes")
            .select("*, resume_items(*), company_analyses(company_name)")
            .eq("id", resume_id)
            .execute()
        )
        if not result.data:
            raise ValueError("자소서를 찾을 수 없습니다.")

        resume = result.data[0]
        items = sorted(resume.get("resume_items", []), key=lambda x: x.get("created_at", ""))
        company = resume.get("company_analyses", {}).get("company_name", "")
        title = resume.get("title", "자기소개서")

        # Build HTML
        html_parts = [
            "<!DOCTYPE html><html><head><meta charset='utf-8'>",
            "<style>",
            "body { font-family: 'Noto Sans KR', sans-serif; max-width: 700px; margin: 40px auto; padding: 0 20px; color: #333; }",
            "h1 { text-align: center; font-size: 22px; margin-bottom: 4px; }",
            ".company { text-align: center; color: #666; font-size: 14px; margin-bottom: 30px; }",
            "h2 { font-size: 14px; color: #1a1a1a; border-bottom: 1px solid #ddd; padding-bottom: 4px; }",
            "p { font-size: 12px; line-height: 1.8; white-space: pre-wrap; }",
            ".meta { font-size: 10px; color: #999; }",
            "</style></head><body>",
            f"<h1>{title}</h1>",
        ]
        if company:
            html_parts.append(f'<div class="company">지원 기업: {company}</div>')

        for i, item in enumerate(items, 1):
            q = item.get("question", "")
            a = item.get("answer", "")
            char_limit = item.get("char_limit")
            html_parts.append(f"<h2>{i}. {q}</h2>")
            html_parts.append(f"<p>{a}</p>")
            if char_limit:
                html_parts.append(f'<div class="meta">글자수: {len(a)}/{char_limit}</div>')

        html_parts.append("</body></html>")
        html_content = "\n".join(html_parts)

        # Try weasyprint, fallback to returning HTML as bytes
        try:
            from weasyprint import HTML
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
        except ImportError:
            # Fallback: return HTML for browser-based PDF printing
            return html_content.encode("utf-8")

    def _load_resume(self, resume_id: str):
        """Load resume data from DB."""
        result = (
            self._db.table("resumes")
            .select("*, resume_items(*), company_analyses(company_name)")
            .eq("id", resume_id)
            .execute()
        )
        if not result.data:
            raise ValueError("자소서를 찾을 수 없습니다.")
        resume = result.data[0]
        items = sorted(resume.get("resume_items", []), key=lambda x: x.get("created_at", ""))
        company = resume.get("company_analyses", {}).get("company_name", "")
        title = resume.get("title", "자기소개서")
        return resume, items, company, title

    def to_md(self, resume_id: str) -> bytes:
        """Generate Markdown from resume data."""
        _, items, company, title = self._load_resume(resume_id)

        lines = [f"# {title}", ""]
        if company:
            lines.append(f"> 지원 기업: {company}")
            lines.append("")

        for i, item in enumerate(items, 1):
            q = item.get("question", "")
            a = item.get("answer", "")
            char_limit = item.get("char_limit")
            lines.append(f"## {i}. {q}")
            lines.append("")
            lines.append(a)
            if char_limit:
                lines.append(f"\n*글자수: {len(a)}/{char_limit}*")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    def to_txt(self, resume_id: str) -> bytes:
        """Generate plain text from resume data."""
        _, items, company, title = self._load_resume(resume_id)

        lines = [title, "=" * len(title.encode("utf-8")), ""]
        if company:
            lines.append(f"지원 기업: {company}")
            lines.append("")

        for i, item in enumerate(items, 1):
            q = item.get("question", "")
            a = item.get("answer", "")
            char_limit = item.get("char_limit")
            lines.append(f"[{i}] {q}")
            lines.append("-" * 40)
            lines.append(a)
            if char_limit:
                lines.append(f"(글자수: {len(a)}/{char_limit})")
            lines.append("")

        return "\n".join(lines).encode("utf-8")

    def to_html(self, resume_id: str) -> bytes:
        """Generate standalone HTML from resume data."""
        _, items, company, title = self._load_resume(resume_id)

        html_parts = [
            "<!DOCTYPE html><html lang='ko'><head><meta charset='utf-8'>",
            f"<title>{title}</title>",
            "<style>",
            "@media print { body { margin: 0; } }",
            "body { font-family: 'Noto Sans KR', -apple-system, sans-serif; max-width: 700px; margin: 40px auto; padding: 0 20px; color: #222; }",
            "h1 { text-align: center; font-size: 22px; margin-bottom: 4px; }",
            ".company { text-align: center; color: #666; font-size: 14px; margin-bottom: 30px; }",
            "h2 { font-size: 14px; color: #1a1a1a; border-bottom: 1px solid #ddd; padding-bottom: 6px; margin-top: 28px; }",
            "p { font-size: 12px; line-height: 1.8; white-space: pre-wrap; }",
            ".meta { font-size: 10px; color: #999; margin-top: 4px; }",
            "</style></head><body>",
            f"<h1>{title}</h1>",
        ]
        if company:
            html_parts.append(f'<div class="company">지원 기업: {company}</div>')

        for i, item in enumerate(items, 1):
            q = item.get("question", "")
            a = item.get("answer", "")
            char_limit = item.get("char_limit")
            html_parts.append(f"<h2>{i}. {q}</h2>")
            html_parts.append(f"<p>{a}</p>")
            if char_limit:
                html_parts.append(f'<div class="meta">글자수: {len(a)}/{char_limit}</div>')

        html_parts.append("</body></html>")
        return "\n".join(html_parts).encode("utf-8")

    def to_json(self, resume_id: str) -> bytes:
        """Generate JSON export of resume data."""
        resume, items, company, title = self._load_resume(resume_id)

        data = {
            "title": title,
            "company": company,
            "status": resume.get("status", ""),
            "created_at": resume.get("created_at", ""),
            "items": [
                {
                    "question": item.get("question", ""),
                    "answer": item.get("answer", ""),
                    "char_limit": item.get("char_limit"),
                    "version": item.get("version", 1),
                    "tone": item.get("tone"),
                }
                for item in items
            ],
        }
        return json_lib.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
